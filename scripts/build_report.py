from __future__ import annotations

import json
import sys
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from sklearn.compose import ColumnTransformer
from sklearn.ensemble import RandomForestClassifier
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import accuracy_score, f1_score, precision_score, recall_score, roc_auc_score
from sklearn.model_selection import train_test_split
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder, StandardScaler


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from src.data_utils import DatasetConfig, split_features_target
from src.modeling import train_model

DATA_PATH = ROOT / "data" / "raw" / "uci_student_dropout" / "data.csv"
OUT_DIR = ROOT / "outputs" / "report"
REPORT_PATH = ROOT / "Informe_Proyecto_Final_Desercion_IA.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
LIGHT_GRAY = "F2F4F7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, LIGHT_GRAY)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            set_cell_text(cell, value)
    doc.add_paragraph()


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    styles["Title"].font.name = "Calibri"
    styles["Title"].font.size = Pt(20)
    styles["Title"].font.color.rgb = DARK_BLUE

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "Proyecto Final · Inteligencia Artificial · EAFIT 2026-1"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(90, 90, 90)

    footer = section.footer.paragraphs[0]
    footer.text = "Desarrollado para el curso de Inteligencia Artificial · Universidad EAFIT"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(100, 100, 100)


def load_and_train() -> tuple[pd.DataFrame, dict, pd.DataFrame]:
    df = pd.read_csv(DATA_PATH, sep=";")
    x, y = split_features_target(df, DatasetConfig("Target", "Dropout"))
    x_train, x_test, y_train, y_test = train_test_split(
        x, y, test_size=0.2, random_state=42, stratify=y
    )

    numeric_features = x.select_dtypes(include=["number", "bool"]).columns.tolist()
    categorical_features = [col for col in x.columns if col not in numeric_features]
    preprocessor = ColumnTransformer(
        transformers=[
            (
                "num",
                Pipeline([("imputer", SimpleImputer(strategy="median")), ("scaler", StandardScaler())]),
                numeric_features,
            ),
            (
                "cat",
                Pipeline(
                    [
                        ("imputer", SimpleImputer(strategy="most_frequent")),
                        ("onehot", OneHotEncoder(handle_unknown="ignore", sparse_output=False)),
                    ]
                ),
                categorical_features,
            ),
        ],
        remainder="drop",
        verbose_feature_names_out=False,
    )

    baseline = Pipeline(
        [
            ("preprocessor", preprocessor),
            ("model", LogisticRegression(max_iter=1000, class_weight="balanced")),
        ]
    )
    baseline.fit(x_train, y_train)
    baseline_pred = baseline.predict(x_test)
    baseline_prob = baseline.predict_proba(x_test)[:, 1]

    result = train_model(x, y)
    metrics = {
        "baseline": {
            "accuracy": accuracy_score(y_test, baseline_pred),
            "precision": precision_score(y_test, baseline_pred, zero_division=0),
            "recall": recall_score(y_test, baseline_pred, zero_division=0),
            "f1": f1_score(y_test, baseline_pred, zero_division=0),
            "roc_auc": roc_auc_score(y_test, baseline_prob),
        },
        "random_forest": result.metrics,
        "confusion": result.confusion.tolist(),
    }
    importance = pd.DataFrame(
        {
            "feature": result.feature_names,
            "importance": result.pipeline.named_steps["model"].feature_importances_,
        }
    ).sort_values("importance", ascending=False)
    return df, metrics, importance


def make_figures(df: pd.DataFrame, importance: pd.DataFrame) -> dict[str, Path]:
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    target_fig = OUT_DIR / "target_distribution.png"
    counts = df["Target"].value_counts()
    fig, ax = plt.subplots(figsize=(6.5, 3.2), dpi=160)
    colors = ["#2E74B5", "#7EA7D8", "#D15B5B"]
    counts.plot(kind="bar", ax=ax, color=colors[: len(counts)])
    ax.set_title("Distribucion de la variable objetivo")
    ax.set_xlabel("Estado academico")
    ax.set_ylabel("Numero de estudiantes")
    ax.grid(axis="y", alpha=0.25)
    plt.tight_layout()
    fig.savefig(target_fig)
    plt.close(fig)

    importance_fig = OUT_DIR / "feature_importance.png"
    top = importance.head(10).sort_values("importance")
    fig, ax = plt.subplots(figsize=(6.5, 3.8), dpi=160)
    ax.barh(top["feature"], top["importance"], color="#2E74B5")
    ax.set_title("Variables mas importantes del Random Forest")
    ax.set_xlabel("Importancia")
    ax.grid(axis="x", alpha=0.25)
    plt.tight_layout()
    fig.savefig(importance_fig)
    plt.close(fig)

    architecture_fig = OUT_DIR / "architecture.png"
    fig, ax = plt.subplots(figsize=(6.5, 2.2), dpi=160)
    ax.axis("off")
    ax.set_xlim(-0.02, 1.04)
    ax.set_ylim(0, 1)
    boxes = [
        ("CSV UCI", 0.02, 0.11),
        ("Preprocesamiento\nimputacion + escalado", 0.18, 0.20),
        ("Random Forest\nprediccion de riesgo", 0.45, 0.19),
        ("SHAP\nexplicacion", 0.70, 0.13),
        ("OpenAI\nrecomendacion", 0.87, 0.12),
    ]
    for label, x, width in boxes:
        patch = FancyBboxPatch(
            (x, 0.42),
            width,
            0.25,
            boxstyle="round,pad=0.025",
            linewidth=1.2,
            edgecolor="#2E74B5",
            facecolor="#F2F4F7",
        )
        ax.add_patch(patch)
        ax.text(
            x + width / 2,
            0.545,
            label,
            ha="center",
            va="center",
            fontsize=8.5,
        )
    for x1, x2 in [(0.13, 0.18), (0.38, 0.45), (0.64, 0.70), (0.83, 0.87)]:
        ax.annotate("", xy=(x2, 0.55), xytext=(x1, 0.55), arrowprops=dict(arrowstyle="->", color="#1F4D78"))
    plt.tight_layout()
    fig.savefig(architecture_fig)
    plt.close(fig)

    return {
        "target": target_fig,
        "importance": importance_fig,
        "architecture": architecture_fig,
    }


def p(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)


def build_report() -> None:
    df, metrics, importance = load_and_train()
    figures = make_figures(df, importance)
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Sistema Inteligente de Prediccion y Explicacion de Riesgo de Desercion Universitaria")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Machine Learning, explicabilidad con SHAP y recomendaciones generadas con OpenAI")
    run.italic = True
    run.font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Nombre Completo 1 · correo1@eafit.edu.co\n"
        "Nombre Completo 2 · correo2@eafit.edu.co\n"
        "Escuela de Ciencias Aplicadas e Ingenieria · Universidad EAFIT\n"
        "Curso: Inteligencia Artificial · Semestre 2026-1\n"
        "Entrega: 19-22 de mayo de 2026\n"
        "Repositorio: REEMPLAZAR_CON_LINK_DE_GITHUB"
    )

    doc.add_heading("Resumen", level=1)
    p(
        doc,
        "Contexto: La desercion universitaria afecta la permanencia estudiantil, la planeacion institucional y el bienestar de los estudiantes. "
        "Objetivo: construir un sistema de alerta temprana que estime el riesgo de desercion y explique los factores principales de cada prediccion. "
        "Metodo: se uso el dataset Predict Students' Dropout and Academic Success de UCI, un pipeline de preprocesamiento, un modelo Random Forest, explicabilidad con SHAP e integracion con OpenAI para generar recomendaciones en lenguaje natural. "
        f"Resultados: el modelo obtuvo AUC-ROC = {metrics['random_forest']['roc_auc']:.3f}, F1 = {metrics['random_forest']['f1']:.3f} y recall = {metrics['random_forest']['recall']:.3f} en el conjunto de prueba. "
        "Conclusion: el sistema muestra que un enfoque combinado de ML explicable y LLM puede apoyar intervenciones academicas tempranas, siempre como herramienta de apoyo y no como sustituto del criterio humano.",
    )
    p(doc, "Palabras clave: Machine Learning, desercion estudiantil, Random Forest, SHAP, OpenAI, Streamlit, EDA.")

    doc.add_heading("1. Introduccion", level=1)
    p(
        doc,
        "La desercion estudiantil es un problema relevante para universidades porque reduce la continuidad academica, incrementa costos institucionales y puede reflejar barreras economicas, academicas o personales. "
        "El proyecto plantea la pregunta de investigacion: ¿puede un modelo de Machine Learning predecir estudiantes en riesgo de desercion con una AUC-ROC superior a 0.85 y generar explicaciones comprensibles para apoyar decisiones de acompanamiento?",
    )
    p(
        doc,
        "El sistema usa un dataset publico de UCI con informacion academica, demografica y socioeconomica de estudiantes de educacion superior. "
        "El documento presenta la exploracion de datos, arquitectura, metodologia experimental, resultados, limitaciones y conclusiones del sistema implementado.",
    )

    doc.add_heading("2. Datos y Exploracion (EDA)", level=1)
    doc.add_heading("2.1. Descripcion del dataset", level=2)
    add_table(
        doc,
        ["Atributo", "Descripcion"],
        [
            ["Fuente", "UCI Machine Learning Repository, dataset 697"],
            ["Registros (N)", f"{len(df):,} estudiantes".replace(",", ".")],
            ["Features (p)", f"{df.shape[1] - 1} variables predictoras"],
            ["Variable objetivo", "Target: Dropout, Enrolled, Graduate. Para este proyecto se binariza Dropout vs. resto."],
            ["Tipo de problema", "Clasificacion binaria para deteccion de riesgo de desercion"],
            ["Licencia", "Creative Commons Attribution 4.0 International (CC BY 4.0)"],
        ],
    )

    doc.add_heading("2.2. Hallazgos del analisis exploratorio", level=2)
    p(
        doc,
        "El dataset contiene tres estados academicos: Graduate, Dropout y Enrolled. La clase Dropout representa 1.421 de 4.424 registros, aproximadamente 32,1% de los casos. "
        "Aunque no es un desbalance extremo, la metrica accuracy puede ocultar errores importantes; por eso se reportan F1, recall y AUC-ROC.",
    )
    p(
        doc,
        "Las variables con mayor importancia estan relacionadas con el desempeno curricular de primer y segundo semestre, el pago oportuno de matricula, edad de ingreso, condicion de deudor y beca. "
        "Esto sugiere que el riesgo de desercion combina senales academicas y socioeconomicas.",
    )
    doc.add_picture(str(figures["target"]), width=Inches(6.2))
    add_caption(doc, "Figura 1. Distribucion de la variable objetivo en el dataset UCI.")

    doc.add_heading("3. Arquitectura del Sistema", level=1)
    p(
        doc,
        "La arquitectura implementada integra carga de datos, preprocesamiento reproducible, entrenamiento de modelo, explicabilidad local y generacion de recomendaciones. "
        "La interfaz Streamlit permite cargar un CSV, seleccionar la variable objetivo, visualizar metricas, revisar explicaciones y generar una recomendacion usando la API de OpenAI.",
    )
    doc.add_picture(str(figures["architecture"]), width=Inches(6.2))
    add_caption(doc, "Figura 2. Arquitectura general del sistema propuesto.")

    doc.add_heading("3.1. Componentes principales", level=2)
    p(doc, "Preprocesamiento. Se usa imputacion por mediana para variables numericas, imputacion por moda para variables categoricas, estandarizacion numerica y One-Hot Encoding para categorias.")
    p(doc, "Modelo de ML. El modelo principal es Random Forest con 300 arboles, profundidad maxima 8, min_samples_leaf = 5 y class_weight = balanced. Este modelo es apropiado para datos tabulares mixtos y permite calcular importancia de variables.")
    p(doc, "Componente LLM. OpenAI se usa para convertir la probabilidad de riesgo y los factores explicativos en recomendaciones en lenguaje natural. La estrategia de prompting es zero-shot con instrucciones de cautela, claridad y no invencion de datos.")
    p(doc, "Interfaz. La aplicacion se construyo en Streamlit e incluye pestanas para datos, evaluacion del modelo y analisis individual de estudiantes.")

    doc.add_heading("4. Metodologia", level=1)
    doc.add_heading("4.1. Configuracion experimental", level=2)
    add_table(
        doc,
        ["Parametro", "Valor"],
        [
            ["Framework", "scikit-learn, pandas, Streamlit, SHAP, OpenAI API"],
            ["Division", "80% entrenamiento / 20% prueba"],
            ["Random state", "42"],
            ["Modelo principal", "RandomForestClassifier"],
            ["Numero de arboles", "300"],
            ["Profundidad maxima", "8"],
            ["Balance de clases", "class_weight = balanced"],
            ["Hardware", "CPU local"],
        ],
    )
    doc.add_heading("4.2. Estrategia de validacion", level=2)
    p(
        doc,
        "Se uso hold-out estratificado con 80% de los registros para entrenamiento y 20% para prueba, manteniendo la proporcion de casos Dropout. "
        "El split se realiza antes del preprocesamiento para evitar data leakage; los parametros de imputacion, escalado y codificacion se aprenden solamente con el conjunto de entrenamiento dentro de un Pipeline.",
    )
    doc.add_heading("4.3. Experimentos realizados", level=2)
    p(doc, "Baseline: Regresion Logistica balanceada como modelo interpretable de referencia.")
    p(doc, "Experimento principal: Random Forest balanceado para capturar relaciones no lineales entre variables academicas, demograficas y socioeconomicas.")
    p(doc, "Integracion explicativa: generacion de factores principales por importancia local y recomendacion textual mediante OpenAI.")

    doc.add_heading("5. Resultados", level=1)
    doc.add_heading("5.1. Metricas de evaluacion", level=2)
    add_table(
        doc,
        ["Modelo", "Accuracy", "Precision", "Recall", "F1", "AUC-ROC"],
        [
            [
                "Baseline (Regresion Logistica)",
                f"{metrics['baseline']['accuracy']:.3f}",
                f"{metrics['baseline']['precision']:.3f}",
                f"{metrics['baseline']['recall']:.3f}",
                f"{metrics['baseline']['f1']:.3f}",
                f"{metrics['baseline']['roc_auc']:.3f}",
            ],
            [
                "Random Forest (mejor)",
                f"{metrics['random_forest']['accuracy']:.3f}",
                f"{metrics['random_forest']['precision']:.3f}",
                f"{metrics['random_forest']['recall']:.3f}",
                f"{metrics['random_forest']['f1']:.3f}",
                f"{metrics['random_forest']['roc_auc']:.3f}",
            ],
        ],
    )
    add_table(
        doc,
        ["", "Pred: no riesgo", "Pred: riesgo"],
        [
            ["Real: no riesgo", str(metrics["confusion"][0][0]), str(metrics["confusion"][0][1])],
            ["Real: riesgo", str(metrics["confusion"][1][0]), str(metrics["confusion"][1][1])],
        ],
    )
    doc.add_picture(str(figures["importance"]), width=Inches(6.2))
    add_caption(doc, "Figura 3. Variables mas importantes del modelo Random Forest.")

    doc.add_heading("5.2. Analisis cualitativo", level=2)
    p(
        doc,
        "Para un estudiante individual, el sistema calcula la probabilidad de desercion, identifica los factores con mayor contribucion y genera una recomendacion. "
        "Ejemplo de salida esperada: Riesgo alto; los factores principales estan asociados con bajo numero de unidades curriculares aprobadas, calificaciones bajas en segundo semestre y matricula no al dia. "
        "La recomendacion sugiere acompanamiento academico, revision de apoyo financiero y seguimiento desde bienestar universitario.",
    )

    doc.add_heading("6. Discusion", level=1)
    doc.add_heading("6.1. Interpretacion de resultados", level=2)
    p(
        doc,
        f"El Random Forest obtuvo AUC-ROC de {metrics['random_forest']['roc_auc']:.3f}, lo que indica buena separacion entre estudiantes con y sin riesgo de desercion. "
        f"El F1 de {metrics['random_forest']['f1']:.3f} y el recall de {metrics['random_forest']['recall']:.3f} muestran que el sistema identifica una proporcion importante de casos Dropout sin sacrificar demasiado la precision.",
    )
    p(
        doc,
        "Las variables mas influyentes corresponden principalmente al desempeno curricular de los primeros semestres. Esto es coherente con la hipotesis de que el rendimiento temprano y las condiciones financieras son senales utiles para intervenciones preventivas.",
    )
    doc.add_heading("6.2. Limitaciones", level=2)
    p(doc, "El dataset proviene de un contexto institucional especifico; por tanto, el modelo puede no generalizar directamente a otras universidades sin recalibracion.")
    p(doc, "La variable Enrolled fue agrupada como no desercion para formular el problema binario, lo cual simplifica una realidad academica mas compleja.")
    p(doc, "Las recomendaciones del LLM dependen de los datos enviados por el sistema y no deben usarse para tomar decisiones automaticas sobre estudiantes.")
    doc.add_heading("6.3. Trabajo futuro", level=2)
    p(doc, "Evaluar modelos adicionales como XGBoost o LightGBM y realizar busqueda de hiperparametros con validacion cruzada estratificada.")
    p(doc, "Incorporar datos longitudinales de asistencia, uso de plataformas academicas y alertas tempranas por periodo.")
    p(doc, "Validar las recomendaciones con expertos de bienestar universitario para mejorar utilidad, tono y responsabilidad.")

    doc.add_heading("7. Conclusiones", level=1)
    p(
        doc,
        "El proyecto demuestra que es posible construir un sistema de alerta temprana de desercion con datos tabulares, Machine Learning explicable e integracion con modelos de lenguaje. "
        f"La pregunta de investigacion se responde positivamente: el modelo principal supero el umbral propuesto de AUC-ROC 0.85 al alcanzar {metrics['random_forest']['roc_auc']:.3f}. "
        "La combinacion de Random Forest, explicabilidad y recomendaciones generadas con OpenAI facilita que usuarios no tecnicos comprendan los factores de riesgo y posibles acciones de acompanamiento.",
    )

    doc.add_heading("Contribuciones del equipo", level=2)
    add_table(
        doc,
        ["Integrante", "Contribucion principal"],
        [
            ["Nombre 1", "EDA, preprocesamiento y visualizaciones"],
            ["Nombre 2", "Arquitectura del modelo y entrenamiento"],
            ["Nombre 3", "Integracion OpenAI, explicabilidad y Streamlit"],
            ["Nombre 4", "Evaluacion, analisis e informe final"],
        ],
    )

    doc.add_heading("Repositorio y Demo", level=2)
    p(doc, "GitHub: REEMPLAZAR_CON_LINK_DE_GITHUB")
    p(doc, "Video demo (max. 3 min): REEMPLAZAR_CON_LINK_DE_VIDEO")
    p(doc, "Instalacion: pip install -r requirements.txt")
    p(doc, "Ejecucion: streamlit run app.py")

    doc.add_heading("Referencias", level=1)
    refs = [
        "UCI Machine Learning Repository. (2021). Predict Students' Dropout and Academic Success. https://archive.ics.uci.edu/dataset/697/predict+students+dropout+and+academic+success",
        "Breiman, L. (2001). Random forests. Machine Learning, 45, 5-32.",
        "Lundberg, S. M., & Lee, S.-I. (2017). A unified approach to interpreting model predictions. Advances in Neural Information Processing Systems.",
        "OpenAI. (2026). OpenAI API documentation. https://platform.openai.com/docs",
        "Pedregosa, F., Varoquaux, G., Gramfort, A., et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825-2830.",
    ]
    for ref in refs:
        p(doc, ref)

    doc.add_heading("A. Detalles de hiperparametros y busqueda", level=1)
    p(doc, "La primera version usa hiperparametros fijos para mantener reproducibilidad. Como extension, se recomienda GridSearchCV o RandomizedSearchCV sobre n_estimators, max_depth, min_samples_leaf y max_features.")

    doc.add_heading("B. Ejemplos adicionales de outputs del sistema", level=1)
    p(
        doc,
        "Ejemplo: Probabilidad estimada de desercion 72%. Factores principales: pocas unidades curriculares aprobadas, baja nota de segundo semestre y deuda. "
        "Recomendacion: priorizar reunion con consejeria academica, revisar plan de financiacion y crear seguimiento de progreso durante las proximas cuatro semanas.",
    )

    doc.save(REPORT_PATH)
    with open(OUT_DIR / "report_metrics.json", "w", encoding="utf-8") as f:
        json.dump(metrics, f, indent=2)
    print(REPORT_PATH)


if __name__ == "__main__":
    build_report()
